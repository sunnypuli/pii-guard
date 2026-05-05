"""
Multi-format file reader and in-place tokenizer.
Supports plain text, CSV, PDF, DOCX, XLSX/XLS.

Optional deps (pip install 'piiwall[rich]'):
  pypdf       — PDF extraction
  python-docx — Word document read/write
  openpyxl    — Excel read/write
"""
from __future__ import annotations

from pathlib import Path

_RICH_EXTS = {".pdf", ".docx", ".xlsx", ".xls"}


def is_rich_format(path: Path) -> bool:
    return path.suffix.lower() in _RICH_EXTS


def read_text(path: Path) -> str:
    """Extract plain text from any supported file format."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in (".xlsx", ".xls"):
        return _read_xlsx(path)
    return path.read_text(encoding="utf-8", errors="replace")


def write_tokenized(
    input_path: Path,
    output_path: Path,
    original_to_token: dict[str, str],
) -> None:
    """
    Write a tokenized copy of input_path to output_path, preserving
    the original file format for DOCX and XLSX. PDF and plain text
    are written as UTF-8 text with the tokenized content.
    """
    suffix = input_path.suffix.lower()
    sorted_map = sorted(original_to_token.items(), key=lambda x: -len(x[0]))

    if suffix == ".docx":
        _tokenize_docx(input_path, output_path, sorted_map)
    elif suffix in (".xlsx", ".xls"):
        _tokenize_xlsx(input_path, output_path, sorted_map)
    else:
        text = read_text(input_path)
        for original, token in sorted_map:
            text = text.replace(original, token)
        output_path.write_text(text, encoding="utf-8")


# ── PDF ───────────────────────────────────────────────────────────────────────

def _read_pdf(path: Path) -> str:
    try:
        import pypdf
    except ImportError:
        raise ImportError(
            "pypdf is required for PDF support.\n"
            "Install it: pip install 'piiwall[rich]'"
        )
    reader = pypdf.PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document support.\n"
            "Install it: pip install 'piiwall[rich]'"
        )
    doc = docx.Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _tokenize_docx(
    input_path: Path,
    output_path: Path,
    sorted_map: list[tuple[str, str]],
) -> None:
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for Word document support.\n"
            "Install it: pip install 'piiwall[rich]'"
        )
    doc = docx.Document(str(input_path))

    for para in doc.paragraphs:
        _replace_in_paragraph(para, sorted_map)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, sorted_map)

    doc.save(str(output_path))


def _replace_in_paragraph(para, sorted_map: list[tuple[str, str]]) -> None:
    full_text = para.text
    new_text = full_text
    for original, token in sorted_map:
        new_text = new_text.replace(original, token)

    if new_text == full_text:
        return

    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(new_text)


# ── XLSX ──────────────────────────────────────────────────────────────────────

def _read_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            "openpyxl is required for Excel support.\n"
            "Install it: pip install 'piiwall[rich]'"
        )
    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            for cell in row:
                if isinstance(cell, str) and cell.strip():
                    parts.append(cell)
    return "\n".join(parts)


def _tokenize_xlsx(
    input_path: Path,
    output_path: Path,
    sorted_map: list[tuple[str, str]],
) -> None:
    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            "openpyxl is required for Excel support.\n"
            "Install it: pip install 'piiwall[rich]'"
        )
    wb = openpyxl.load_workbook(str(input_path))
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    new_val = cell.value
                    for original, token in sorted_map:
                        new_val = new_val.replace(original, token)
                    cell.value = new_val
    wb.save(str(output_path))
